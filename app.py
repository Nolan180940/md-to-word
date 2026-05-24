import streamlit as st
import pypandoc
import tempfile
import os
import re
import secrets
from datetime import datetime
from typing import Optional, List, Tuple

# 尝试导入 python-docx
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# 常量定义
MAX_CONTENT_LENGTH = 100000
SESSION_HISTORY_KEY = "edit_history"
MAX_HISTORY_COUNT = 5

# 主题模板
THEME_TEMPLATES = {
    "学术论文": {"font": "Times New Roman", "heading_font": "Arial", "code_font": "Consolas", "line_spacing": 1.5, "heading_colors": ["#000000", "#2E5C8A", "#4A7BA7"], "code_bg": "F5F5F5", "include_cover": True, "include_toc": True, "include_header_footer": True},
    "商务报告": {"font": "Calibri", "heading_font": "Calibri Light", "code_font": "Consolas", "line_spacing": 1.15, "heading_colors": ["#1F4E79", "#2E75B6", "#5B9BD5"], "code_bg": "F0F0F0", "include_cover": True, "include_toc": True, "include_header_footer": True},
    "技术文档": {"font": "Segoe UI", "heading_font": "Segoe UI Light", "code_font": "Consolas", "line_spacing": 1.2, "heading_colors": ["#2B579A", "#4472C4", "#6C9EEB"], "code_bg": "F8F8F8", "include_cover": False, "include_toc": True, "include_header_footer": False},
    "简约风格": {"font": "Arial", "heading_font": "Arial", "code_font": "Courier New", "line_spacing": 1.0, "heading_colors": ["#000000", "#333333", "#666666"], "code_bg": "FAFAFA", "include_cover": False, "include_toc": False, "include_header_footer": False},
}

st.set_page_config(page_title="Markdown to Word Pro", page_icon="🎨", layout="wide")

def sanitize_filename(raw_title: str) -> str:
    if not raw_title or not raw_title.strip():
        return "document"
    clean_name = re.sub(r'[^a-zA-Z0-9_\u4e00-\u9fff\s-]', '', raw_title)
    clean_name = re.sub(r'\s+', '-', clean_name).strip('-_')
    return (clean_name[:50] or "document")

def generate_secure_filename(text: str) -> str:
    h1_match = re.search(r'^#\s+(.+)$', text, re.MULTILINE)
    raw_title = h1_match.group(1).strip() if h1_match else "document"
    clean_name = sanitize_filename(raw_title)
    random_suffix = secrets.token_hex(4)
    return f"{clean_name}_{random_suffix}.docx"

def extract_headings(text: str) -> List[Tuple[int, str]]:
    headings = []
    for line in text.split('\n'):
        match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if match:
            level = len(match.group(1))
            title = re.sub(r'\$\$?.*?\$\$?', '', match.group(2)).strip()
            if title:
                headings.append((level, title))
    return headings

def save_to_history(text: str):
    if SESSION_HISTORY_KEY not in st.session_state:
        st.session_state[SESSION_HISTORY_KEY] = []
    history = st.session_state[SESSION_HISTORY_KEY]
    timestamp = datetime.now().strftime("%H:%M:%S")
    if history and history[-1][1] == text:
        return
    history.append((timestamp, text))
    if len(history) > MAX_HISTORY_COUNT:
        history.pop(0)
    st.session_state[SESSION_HISTORY_KEY] = history

def get_history():
    return st.session_state.get(SESSION_HISTORY_KEY, [])

def rollback_to_version(index: int) -> Optional[str]:
    history = get_history()
    return history[index][1] if 0 <= index < len(history) else None

def smart_fix_markdown(text: str) -> Tuple[str, List[str]]:
    if not text:
        return "", []
    log = []
    fixed_text = text
    original_text = text
    
    # 清理零宽字符
    for char in ['\u200b', '\u200c', '\u200d', '\ufeff']:
        if char in fixed_text:
            fixed_text = fixed_text.replace(char, '')
            log.append(f"🧹 移除了隐形字符")
    
    # 标准化 LaTeX 公式
    replacements = [('\\[', '$$'), ('\\]', '$$'), ('\\(', '$'), ('\\)', '$')]
    for old, new in replacements:
        if old in fixed_text:
            fixed_text = fixed_text.replace(old, new)
    if '\\[' in text or '\\(' in text:
        log.append("📐 标准化了 LaTeX 公式语法")
    
    # 修复公式空格
    pattern = r'(?<!\$)\$[ \t]+([^ \t][^$]*?[^ \t])[ \t]+\$(?!\$)'
    matches = re.findall(pattern, fixed_text)
    if matches:
        fixed_text = re.sub(pattern, r'$\1$', fixed_text)
        log.append(f"🔧 修复了 {len(matches)} 处行内公式空格")
    
    # HTML 标签转换
    html_conv = [(r'<sup>(.*?)</sup>', r'^\1^', '上标'), (r'<sub>(.*?)</sub>', r'~\1~', '下标')]
    for pattern, repl, desc in html_conv:
        matches = re.findall(pattern, fixed_text, re.IGNORECASE)
        if matches:
            fixed_text = re.sub(pattern, repl, fixed_text, flags=re.IGNORECASE)
            log.append(f"🔄 转换了 {len(matches)} 处 HTML {desc}标签")
    
    # 闭合代码块和公式块
    if len(re.findall(r'^```', fixed_text, re.MULTILINE)) % 2 != 0:
        fixed_text += "\n```\n"
        log.append("🧱 自动闭合了代码块")
    if fixed_text.count('$$') % 2 != 0:
        fixed_text += "\n$$\n"
        log.append("🧮 自动闭合了公式块")
    
    # 格式化优化
    fixed_text = re.sub(r'([^\n])\n(\s*```)', r'\1\n\n\2', fixed_text)
    fixed_text = re.sub(r'(```\s*)\n([^\n])', r'\1\n\n\2', fixed_text)
    
    # Blockquote 处理
    lines = fixed_text.splitlines()
    if lines:
        out_lines, i, made_change = [], 0, False
        while i < len(lines):
            stripped = lines[i].lstrip()
            if stripped.startswith('>'):
                if out_lines and out_lines[-1].strip():
                    out_lines.append('')
                    made_change = True
                while i < len(lines) and lines[i].lstrip().startswith('>'):
                    out_lines.append(lines[i])
                    i += 1
                if i < len(lines) and lines[i].strip():
                    out_lines.append('')
                    made_change = True
                continue
            out_lines.append(lines[i])
            i += 1
        if made_change:
            fixed_text = '\n'.join(out_lines)
            log.append("🧩 优化了引用块段落间距")
    
    if fixed_text == original_text:
        log.append("✅ 无需修复，格式良好")
    return fixed_text, log

def apply_word_styles(docx_path: str, theme: str = "学术论文"):
    if not HAS_DOCX:
        return
    theme_config = THEME_TEMPLATES.get(theme, THEME_TEMPLATES["学术论文"])
    try:
        doc = Document(docx_path)
        styles = doc.styles
        default_style = styles['Normal']
        default_style.font.name = theme_config["font"]
        
        # 代码块样式
        style_name = 'Source Code' if 'Source Code' in styles else 'SourceCode'
        if style_name in styles:
            style_code = styles[style_name]
            style_code.font.name = theme_config["code_font"]
            p_pr = style_code.element.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:fill'), theme_config["code_bg"])
            p_pr.append(shd)
        
        # 标题样式
        for i, color in enumerate(theme_config["heading_colors"][:3]):
            heading_name = f'Heading {i+1}'
            if heading_name in styles:
                heading_style = styles[heading_name]
                heading_style.font.name = theme_config["heading_font"]
                heading_style.font.color.rgb = RGBColor(int(color[1:3], 16), int(color[3:5], 16), int(color[5:7], 16))
        doc.save(docx_path)
    except Exception as e:
        print(f"样式应用失败：{e}")

def add_cover_page(doc, title: str, theme: str):
    theme_config = THEME_TEMPLATES.get(theme, THEME_TEMPLATES["学术论文"])
    cover_para = doc.add_paragraph()
    cover_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover_para.add_run(title)
    run.font.name = theme_config["heading_font"]
    run.font.size = Pt(24)
    run.bold = True
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"\n\n{datetime.now().strftime('%Y年%m月%d日')}")
    date_run.font.size = Pt(12)
    doc.add_page_break()

def add_header_footer(doc, theme: str):
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_run = header_para.add_run("内部资料 • 保密")
    header_run.font.size = Pt(9)
    header_run.font.color.rgb = RGBColor(128, 128, 128)
    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run("第 { PAGE } 页")
    footer_run.font.size = Pt(9)

def convert_to_docx(md_content: str, theme: str = "学术论文") -> Tuple[Optional[str], Optional[str]]:
    if len(md_content) > MAX_CONTENT_LENGTH:
        return None, f"内容过长（{len(md_content)} 字符）"
    output_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
            output_path = tmp_file.name
        processed_content = md_content
        mermaid_blocks = re.findall(r'```mermaid\s*(.*?)\s*```', md_content, re.DOTALL)
        for i, _ in enumerate(mermaid_blocks):
            processed_content = re.sub(r'```mermaid.*?```', f'![Mermaid 图表 {i+1}](mermaid_{i+1}.png)', processed_content, count=1, flags=re.DOTALL)
        pypandoc.convert_text(processed_content, 'docx', format='markdown+tex_math_dollars', outputfile=output_path, extra_args=['--standalone', '--toc'])
        if HAS_DOCX:
            apply_word_styles(output_path, theme)
            theme_config = THEME_TEMPLATES.get(theme, THEME_TEMPLATES["学术论文"])
            doc = Document(output_path)
            headings = extract_headings(md_content)
            main_title = headings[0][1] if headings else "文档"
            if theme_config.get("include_cover", False):
                temp_doc = Document()
                add_cover_page(temp_doc, main_title, theme)
                for para in doc.paragraphs:
                    temp_doc.add_paragraph(para.text)
                temp_doc.save(output_path)
                doc = Document(output_path)
            if theme_config.get("include_header_footer", False):
                add_header_footer(doc, theme)
                doc.save(output_path)
        return output_path, None
    except Exception as e:
        if output_path and os.path.exists(output_path):
            try:
                os.remove(output_path)
            except:
                pass
        return None, str(e)

# 初始化 session state
if 'current_text' not in st.session_state:
    st.session_state.current_text = ""
if 'selected_theme' not in st.session_state:
    st.session_state.selected_theme = "学术论文"

st.title("🛠️ Markdown 转 Word 专业版")
st.caption("🎨 主题定制 | 📑 自动目录 | 🖼️ 图表支持 | 📜 版本历史 | 🔒 安全增强")
st.divider()

if not HAS_DOCX:
    st.error("⚠️ 未安装 python-docx，样式增强功能将无法生效。")

with st.sidebar:
    st.header("⚙️ 设置")
    selected_theme = st.selectbox("文档风格", options=list(THEME_TEMPLATES.keys()), index=0)
    st.session_state.selected_theme = selected_theme
    theme_info = THEME_TEMPLATES[selected_theme]
    with st.expander("📋 主题详情"):
        st.write(f"**字体**: {theme_info['font']}")
        st.write(f"**封面**: {'✓' if theme_info['include_cover'] else '✗'}")
        st.write(f"**目录**: {'✓' if theme_info['include_toc'] else '✗'}")
    st.divider()
    st.subheader("📜 版本历史")
    history = get_history()
    if history:
        for i, (timestamp, content) in enumerate(reversed(history)):
            preview = content[:40].replace('\n', ' ') + "..."
            if st.button(f"⏪ {timestamp} - {preview}", key=f"history_{i}"):
                st.session_state.current_text = rollback_to_version(len(history) - 1 - i)
                st.rerun()
    else:
        st.info("暂无历史记录")
    with st.expander("❓ 使用说明"):
        st.markdown("**功能**: LaTeX 公式修复、代码高亮、Mermaid 图表、自动目录、版本历史\n\n**安全**: 文件名白名单过滤、内容长度限制、临时文件自动清理")

default_text = r'''# 深度学习中的概率分布

## 1. 公式测试

行内公式： $E = mc^2$ 和 $ x_0 = 0 $ 会自动修复为 $E=mc^2$ 和 $x_0=0$。

块级公式：
\[
\mathcal{L}(\theta) = -\frac{1}{N} \sum_{i=1}^N y_i \log(\hat{y}_i)
\]

## 2. 代码块

```python
def hello():
    print("Hello World")
```

## 3. Mermaid 图表

```mermaid
graph TD
    A[开始] --> B{判断}
    B -->|是 | C[执行 A]
    B -->|否 | D[执行 B]
```

## 4. 引用块

> 这是一段引用
> 可以有多行
'''

col_input, col_preview = st.columns(2, gap="medium")

with col_input:
    st.subheader("⌨️ 编辑区")
    md_text = st.text_area("Input", value=st.session_state.get('current_text', default_text), height=600, label_visibility="collapsed", placeholder="在此粘贴 Markdown...", key="md_editor")
    st.session_state.current_text = md_text
    col_btn1, col_btn2, col_btn3 = st.columns(3)
    with col_btn1:
        if st.button("📋 清空", use_container_width=True):
            st.session_state.current_text = ""
            st.rerun()
    with col_btn2:
        if st.button("💾 保存快照", use_container_width=True) and md_text.strip():
            save_to_history(md_text)
            st.success("已保存")
    with col_btn3:
        if st.button("📥 示例", use_container_width=True):
            st.session_state.current_text = default_text
            st.rerun()

with col_preview:
    st.subheader("👁️ 实时预览")
    preview_text, logs = smart_fix_markdown(md_text)
    if logs and any("无需修复" not in log for log in logs):
        with st.expander(f"🤖 自动修复 ({len([l for l in logs if '无需修复' not in l])} 项)", expanded=True):
            for log in logs:
                if "无需修复" not in log:
                    st.markdown(f"- {log}")
    with st.container(border=True):
        if preview_text.strip():
            st.markdown(preview_text)
            c1, c2, c3 = st.columns(3)
            c1.metric("字符数", len(preview_text))
            c2.metric("标题数", len(extract_headings(preview_text)))
            c3.metric("代码块", len(re.findall(r'```', preview_text)) // 2)
        else:
            st.write("等待输入...")

st.divider()
st.subheader("🚀 生成选项")
auto_clean = st.checkbox("🧹 完成后清理缓存", value=True)

col1, col2, col3 = st.columns([1, 2, 1])
with col2:
    if st.button("🚀 生成定制化 Word 文档", type="primary", use_container_width=True):
        if not md_text.strip():
            st.warning("⚠️ 内容不能为空")
        else:
            save_to_history(md_text)
            final_text, _ = smart_fix_markdown(md_text)
            file_name = generate_secure_filename(final_text)
            with st.spinner("正在渲染..."):
                docx_path, error_msg = convert_to_docx(final_text, theme=selected_theme)
            if docx_path:
                with open(docx_path, "rb") as f:
                    file_data = f.read()
                st.success(f"✅ 生成成功：**{file_name}**")
                st.download_button(label="⬇️ 点击下载 Word", data=file_data, file_name=file_name, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
                if auto_clean:
                    try:
                        os.remove(docx_path)
                        st.info("🧹 临时文件已清理")
                    except:
                        pass
            else:
                st.error(f"❌ 转换失败：{error_msg}")

st.divider()
st.caption("💡 提示：生成的 Word 包含自动目录，可在 Word 中右键更新域刷新页码")
